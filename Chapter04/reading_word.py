import docx

doc = docx.Document('documents/document-1.docx')

# for index, paragraphs in enumerate(doc.paragraphs):
#     if paragraphs.text:
#         print(index, paragraphs.text)

pr_runs = [run.text for run in doc.paragraphs[48].runs]
print(pr_runs)