from docx.shared import Inches, Cm
import docx

document = docx.Document()
p = document.add_paragraph('This is the start of the paragraph')
section = document.add_section( docx.enum.section.WD_SECTION.ODD_PAGE)
section.page_height = Inches(10)
section.page_width = Cm(8)
section.left_margin = Cm(1)
section.right_margin = Cm(1)
section.top_margin = Inches(1)
section.bottom_margin = Cm(1)
document.add_paragraph('This is the start of the paragraph')

document.save('custom-size-document.docx')