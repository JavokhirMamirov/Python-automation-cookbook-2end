import docx
from datetime import datetime

constext = {
    'date': datetime.now(),
    'movies': ['Casablanca', 'The Sound of Music', 'Vertigo'],
    'total_minutes': 404,
}

document = docx.Document()
document.add_heading('Movies Report', 0)

paragraph = document.add_paragraph('Date: ')
paragraph.add_run(str(constext['date'])).italic = True

paragraph = document.add_paragraph('Movies see in the last 30 days: ')
paragraph.add_run(str(len(constext['movies']))).italic = True

for movie in constext['movies']:
    document.add_paragraph(movie, style="List Bullet")

paragraph = document.add_paragraph('Total minutes: ')
paragraph.add_run(str(constext['total_minutes'])).italic = True

document.save('word-report.docx')