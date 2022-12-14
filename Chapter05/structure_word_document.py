import docx

document = docx.Document()

p = document.add_paragraph('This is the start of the paragraph')
run = p.add_run()
run.add_break(docx.text.run.WD_BREAK.LINE)
p.add_run('And now this in a differnt line')
p.add_run(". Even if it's on the same paragraph")
document.add_page_break()
document.add_paragraph('This appears in a new page')

section = document.add_section(docx.enum.section.WD_SECTION.NEW_PAGE)
section.orientarion = docx.enum.section.WD_ORIENT.LANDSCAPE
section.page_height, section.page_width = section.page_width, section.page_height
document.add_paragraph('This is part of a new landscape section')

section = document.add_section(docx.enum.section.WD_SECTION.NEW_PAGE)
section.orientation = docx.enum.section.WD_ORIENT.PORTRAIT
section.page_height, section.page_width = section.page_width, section.page_height
document.add_paragraph('In this section, recover the portrait orientation')
document.save('word-report-structure.docx')
